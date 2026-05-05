"""app/server/pi_ceo_docparser.py — RA-1995 (Wave 2 / 6).

Document parser for ICP research input — customer interview transcripts,
market reports, regulatory docs. Pure deterministic extraction (no LLM
call); the marketing-icp-research / Margot Wave 5 research skills
consume the parsed output.

Public API:
    parse_document(path, *, max_pages=None) -> ParsedDoc

Supported types (auto-detected by file extension):
  * `.pdf`  — PyMuPDF preferred, pypdf fallback, `_(missing_dep)_` if neither
  * `.docx` — python-docx
  * `.txt`  — pure stdlib

Returned ``ParsedDoc`` carries:
  * ``text`` — body text, joined per page with form-feed separators so
    downstream readers can recover page boundaries.
  * ``pages`` — list of (page_number_1_indexed, page_text) tuples.
  * ``tables`` — list of list-of-lists (one entry per detected table,
    DOCX only — PDF table extraction is intentionally out of scope for
    Wave 2; PyMuPDF's table heuristics are unreliable enough that
    surfacing them as ``tables`` would mislead consumers).
  * ``metadata`` — title / author / num_pages / source_path.
  * ``error`` — populated only when parsing failed catastrophically.
    Missing optional deps return a structured error rather than raising.

Fail-soft on missing deps:
  PyMuPDF (`pymupdf`) is preferred for PDFs but is licensed AGPL —
  some deployments will skip it. The module falls through to `pypdf`
  (BSD, pure Python). If both are missing, returns
  `error="missing_pdf_dep"` with `fix` hinting `pip install pymupdf`.

Path safety:
  Caller is responsible for sandboxing — this module reads any path
  it's given. Do NOT expose this directly to user input from the web
  surface; use it from server-side code paths that have already
  validated the source.

Refs: RA-1988 (Wave 2 epic), RA-1971 audit, marketing-icp-research SKILL.
"""
from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path

log = logging.getLogger("pi-ceo.docparser")

# Tunables
DEFAULT_MAX_PAGES: int = 500


@dataclass
class ParsedDoc:
    """Output of `parse_document()`."""

    text: str = ""
    pages: list[tuple[int, str]] = field(default_factory=list)
    tables: list[list[list[str]]] = field(default_factory=list)
    metadata: dict[str, str | int] = field(default_factory=dict)
    source_path: str = ""
    file_type: str = ""  # 'pdf' / 'docx' / 'txt' / 'unknown'
    error: str | None = None

    @property
    def num_pages(self) -> int:
        return len(self.pages) or int(self.metadata.get("num_pages", 0) or 0)

    @property
    def num_chars(self) -> int:
        return len(self.text)


# ── Type detection ───────────────────────────────────────────────────────────


def _detect_type(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return "pdf"
    if suffix == ".docx":
        return "docx"
    if suffix in (".txt", ".md"):
        return "txt"
    return "unknown"


# ── PDF parsing ──────────────────────────────────────────────────────────────


def _parse_pdf(path: Path, *, max_pages: int) -> ParsedDoc:
    """Try PyMuPDF first (higher text fidelity, faster). Fall back to
    pypdf (pure-Python, BSD). Return structured error if neither
    available."""
    out = ParsedDoc(source_path=str(path), file_type="pdf")
    # Try PyMuPDF — preferred for fidelity
    try:
        import pymupdf  # type: ignore  # noqa: PLC0415

        try:
            doc = pymupdf.open(str(path))
        except Exception as exc:  # noqa: BLE001
            out.error = f"pymupdf_open_failed: {exc}"
            return out
        pages: list[tuple[int, str]] = []
        for i, page in enumerate(doc, start=1):
            if i > max_pages:
                break
            try:
                pages.append((i, page.get_text("text") or ""))
            except Exception as exc:  # noqa: BLE001
                pages.append((i, f"_(page_{i}_extract_failed: {exc})_"))
        out.pages = pages
        out.text = "\f".join(p for _, p in pages)
        meta = doc.metadata or {}
        out.metadata = {
            "title": meta.get("title", "") or "",
            "author": meta.get("author", "") or "",
            "num_pages": doc.page_count,
            "source_path": str(path),
        }
        doc.close()
        return out
    except ImportError:
        pass

    # Fallback — pypdf
    try:
        import pypdf  # type: ignore  # noqa: PLC0415

        try:
            reader = pypdf.PdfReader(str(path))
        except Exception as exc:  # noqa: BLE001
            out.error = f"pypdf_open_failed: {exc}"
            return out
        pages2: list[tuple[int, str]] = []
        for i, page in enumerate(reader.pages, start=1):
            if i > max_pages:
                break
            try:
                pages2.append((i, page.extract_text() or ""))
            except Exception as exc:  # noqa: BLE001
                pages2.append((i, f"_(page_{i}_extract_failed: {exc})_"))
        out.pages = pages2
        out.text = "\f".join(p for _, p in pages2)
        info = getattr(reader, "metadata", None) or {}
        out.metadata = {
            "title": str(info.get("/Title", "") or "") if info else "",
            "author": str(info.get("/Author", "") or "") if info else "",
            "num_pages": len(reader.pages),
            "source_path": str(path),
        }
        return out
    except ImportError:
        pass

    out.error = "missing_pdf_dep"
    out.metadata = {
        "fix": "pip install pymupdf  # OR  pip install pypdf",
        "source_path": str(path),
    }
    return out


# ── DOCX parsing ─────────────────────────────────────────────────────────────


def _parse_docx(path: Path, *, max_pages: int) -> ParsedDoc:
    """python-docx is the de-facto DOCX parser. Pages aren't a native
    concept in DOCX; we return one synthetic page per logical section
    (page break or top-of-doc). Tables are extracted as list-of-rows."""
    out = ParsedDoc(source_path=str(path), file_type="docx")
    try:
        import docx  # type: ignore  # noqa: PLC0415
    except ImportError:
        out.error = "missing_docx_dep"
        out.metadata = {
            "fix": "pip install python-docx",
            "source_path": str(path),
        }
        return out
    try:
        doc = docx.Document(str(path))
    except Exception as exc:  # noqa: BLE001
        out.error = f"docx_open_failed: {exc}"
        return out

    # Walk paragraphs, splitting at hard page breaks.
    page_buf: list[str] = []
    pages: list[tuple[int, str]] = []
    page_num = 1
    for para in doc.paragraphs:
        text = para.text or ""
        page_buf.append(text)
        # Check for hard page-break runs
        for run in para.runs:
            # python-docx exposes page-break via the underlying XML;
            # checking the element XML keeps this independent of
            # internal API churn.
            xml = (run._element.xml if hasattr(run, "_element") else "") or ""
            if 'w:type="page"' in xml:
                pages.append((page_num, "\n".join(page_buf).strip()))
                page_num += 1
                page_buf = []
                if page_num > max_pages:
                    break
        if page_num > max_pages:
            break
    if page_buf and page_num <= max_pages:
        pages.append((page_num, "\n".join(page_buf).strip()))

    # Tables
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        tables.append(rows)

    # Metadata via core_properties
    cp = getattr(doc, "core_properties", None)
    out.metadata = {
        "title": (cp.title if cp and cp.title else "") if cp else "",
        "author": (cp.author if cp and cp.author else "") if cp else "",
        "num_pages": len(pages),
        "source_path": str(path),
    }
    out.pages = pages
    out.text = "\f".join(p for _, p in pages)
    out.tables = tables
    return out


# ── TXT parsing ──────────────────────────────────────────────────────────────


def _parse_txt(path: Path, *, max_pages: int) -> ParsedDoc:
    """Read a plain-text file. Synthetic single-page output."""
    out = ParsedDoc(source_path=str(path), file_type="txt")
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
    except OSError as exc:
        out.error = f"read_failed: {exc}"
        return out
    out.text = text
    out.pages = [(1, text)]
    out.metadata = {
        "title": path.stem,
        "author": "",
        "num_pages": 1,
        "source_path": str(path),
    }
    # max_pages doesn't really apply to txt — single-page output
    _ = max_pages
    return out


# ── Public API ───────────────────────────────────────────────────────────────


def parse_document(path: str | Path,
                    *,
                    max_pages: int | None = None) -> ParsedDoc:
    """Parse a PDF / DOCX / TXT into a structured ``ParsedDoc``.

    Args:
        path: Local filesystem path. The caller is responsible for
            sandboxing the source; this module reads whatever it's given.
        max_pages: Cap pages parsed (default 500). Cheap protection
            against unbounded extraction during ICP research workflows.

    Returns:
        ``ParsedDoc`` with ``error`` populated on any failure mode
        (missing dep, unreadable file, unknown extension). Never raises
        — research workflows should be able to inspect ``error`` and
        skip, not crash the pipeline.
    """
    p = Path(path)
    cap = int(max_pages) if max_pages and max_pages > 0 else DEFAULT_MAX_PAGES
    if not p.is_file():
        return ParsedDoc(
            source_path=str(p), file_type="unknown",
            error=f"file_not_found: {p}",
        )
    file_type = _detect_type(p)
    if file_type == "pdf":
        return _parse_pdf(p, max_pages=cap)
    if file_type == "docx":
        return _parse_docx(p, max_pages=cap)
    if file_type == "txt":
        return _parse_txt(p, max_pages=cap)
    return ParsedDoc(
        source_path=str(p), file_type="unknown",
        error=f"unsupported_extension: {p.suffix or '(none)'}",
        metadata={"source_path": str(p)},
    )


__all__ = [
    "ParsedDoc",
    "parse_document",
    "DEFAULT_MAX_PAGES",
]
