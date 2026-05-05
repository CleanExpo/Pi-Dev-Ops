"""tests/test_pi_ceo_docparser.py — RA-1995 regression coverage.

Covers:
  * TXT parsing (always available, pure stdlib)
  * DOCX parsing — generated synthetically via python-docx, so fixture
    creation is part of the test (no checked-in binary blobs)
  * PDF parsing — synthetic PDF only when pymupdf or pypdf is available;
    otherwise asserts the structured `missing_pdf_dep` error
  * Error paths: file_not_found, unsupported extension, missing dep
  * `parse_document()` never raises — all failures land on `ParsedDoc.error`
"""
from __future__ import annotations

import sys
from pathlib import Path

import pytest

REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT))

from app.server.pi_ceo_docparser import (  # noqa: E402
    DEFAULT_MAX_PAGES,
    ParsedDoc,
    parse_document,
)


def _have_pdf_lib() -> bool:
    try:
        import pymupdf  # noqa: F401
        return True
    except ImportError:
        try:
            import pypdf  # noqa: F401
            return True
        except ImportError:
            return False


def _have_docx_lib() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False


# ── TXT ──────────────────────────────────────────────────────────────────────


def test_txt_parsed_to_single_page(tmp_path):
    p = tmp_path / "notes.txt"
    p.write_text("first line\nsecond line\n")
    doc = parse_document(p)
    assert doc.error is None
    assert doc.file_type == "txt"
    assert doc.num_pages == 1
    assert "first line" in doc.text
    assert doc.metadata["title"] == "notes"
    assert doc.num_chars > 0


def test_md_treated_as_txt(tmp_path):
    p = tmp_path / "notes.md"
    p.write_text("# heading\n\nbody")
    doc = parse_document(p)
    assert doc.error is None
    assert doc.file_type == "txt"
    assert "heading" in doc.text


def test_txt_unicode_safe(tmp_path):
    p = tmp_path / "notes.txt"
    p.write_text("éàü 你好 🚀", encoding="utf-8")
    doc = parse_document(p)
    assert doc.error is None
    assert "éàü" in doc.text
    assert "你好" in doc.text


# ── Path validation ──────────────────────────────────────────────────────────


def test_missing_file_returns_error(tmp_path):
    doc = parse_document(tmp_path / "nope.pdf")
    assert doc.error is not None
    assert "file_not_found" in doc.error


def test_unsupported_extension_returns_error(tmp_path):
    p = tmp_path / "data.xlsx"
    p.write_bytes(b"PK\x03\x04 stub")
    doc = parse_document(p)
    assert doc.error is not None
    assert "unsupported_extension" in doc.error
    assert ".xlsx" in doc.error


def test_no_extension_returns_error(tmp_path):
    p = tmp_path / "stub"
    p.write_text("plain")
    doc = parse_document(p)
    assert doc.error is not None
    assert "unsupported_extension" in doc.error


# ── DOCX ─────────────────────────────────────────────────────────────────────


@pytest.mark.skipif(not _have_docx_lib(), reason="python-docx not installed")
def test_docx_parsed_with_paragraphs(tmp_path):
    import docx  # type: ignore

    src = tmp_path / "interview.docx"
    d = docx.Document()
    d.core_properties.title = "Customer interview"
    d.core_properties.author = "Phill"
    d.add_paragraph("Pain point one.")
    d.add_paragraph("Pain point two.")
    d.save(str(src))

    doc = parse_document(src)
    assert doc.error is None
    assert doc.file_type == "docx"
    assert "Pain point one." in doc.text
    assert "Pain point two." in doc.text
    assert doc.metadata["title"] == "Customer interview"
    assert doc.metadata["author"] == "Phill"


@pytest.mark.skipif(not _have_docx_lib(), reason="python-docx not installed")
def test_docx_extracts_tables(tmp_path):
    import docx  # type: ignore

    src = tmp_path / "report.docx"
    d = docx.Document()
    d.add_paragraph("Pricing table:")
    table = d.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Tier"
    table.cell(0, 1).text = "Monthly"
    table.cell(1, 0).text = "Solo"
    table.cell(1, 1).text = "$29"
    table.cell(2, 0).text = "Team"
    table.cell(2, 1).text = "$99"
    d.save(str(src))

    doc = parse_document(src)
    assert doc.error is None
    assert len(doc.tables) == 1
    rows = doc.tables[0]
    assert rows[0] == ["Tier", "Monthly"]
    assert rows[1] == ["Solo", "$29"]
    assert rows[2] == ["Team", "$99"]


# ── PDF ──────────────────────────────────────────────────────────────────────


def test_pdf_missing_dep_returns_structured_error(tmp_path):
    """If neither pymupdf nor pypdf is installed, parse_document
    returns ParsedDoc(error='missing_pdf_dep') with a fix hint —
    NEVER raises."""
    src = tmp_path / "report.pdf"
    src.write_bytes(b"%PDF-1.4 stub\n")
    doc = parse_document(src)
    if _have_pdf_lib():
        # When a PDF lib IS installed, this stub will fail to parse —
        # but the error must be a structured `*_open_failed`, not a
        # bare exception, and `error` must be populated.
        assert doc.error is not None
        assert ("open_failed" in doc.error
                or "missing_pdf_dep" in doc.error)
    else:
        assert doc.error == "missing_pdf_dep"
        assert "fix" in doc.metadata


@pytest.mark.skipif(not _have_pdf_lib(), reason="no PDF lib installed")
def test_pdf_real_synthetic_extracts_text(tmp_path):
    """Build a tiny PDF with a known string and confirm it round-trips."""
    src = tmp_path / "tiny.pdf"
    # Build a minimal PDF using whichever lib is available
    try:
        import pymupdf  # type: ignore

        d = pymupdf.open()
        page = d.new_page()
        page.insert_text((72, 72), "Hello docparser RA-1995")
        d.save(str(src))
        d.close()
    except ImportError:
        # Fallback — pypdf can READ PDFs but not write them. We can use
        # `reportlab` if it's around, otherwise skip this test on the
        # pypdf-only path.
        try:
            from reportlab.pdfgen import canvas  # type: ignore

            c = canvas.Canvas(str(src))
            c.drawString(72, 720, "Hello docparser RA-1995")
            c.save()
        except ImportError:
            pytest.skip("reportlab not available and pymupdf-only path")
    doc = parse_document(src)
    assert doc.error is None
    assert doc.file_type == "pdf"
    assert "Hello docparser RA-1995" in doc.text


# ── max_pages cap ────────────────────────────────────────────────────────────


def test_default_max_pages_constant():
    assert DEFAULT_MAX_PAGES == 500


@pytest.mark.skipif(not _have_docx_lib(), reason="python-docx not installed")
def test_docx_max_pages_caps_extraction(tmp_path):
    """Synthesise a multi-page DOCX and verify max_pages caps output."""
    import docx  # type: ignore

    src = tmp_path / "long.docx"
    d = docx.Document()
    for i in range(5):
        para = d.add_paragraph(f"page {i + 1} content")
        # Insert a hard page break after each paragraph except last
        if i < 4:
            run = para.add_run()
            run.add_break(docx.enum.text.WD_BREAK.PAGE)
    d.save(str(src))

    doc_full = parse_document(src)
    assert doc_full.error is None
    full_pages = doc_full.num_pages

    doc_capped = parse_document(src, max_pages=2)
    assert doc_capped.error is None
    assert doc_capped.num_pages <= 2
    assert doc_capped.num_pages < full_pages


# ── ParsedDoc surface ────────────────────────────────────────────────────────


def test_parsed_doc_defaults():
    pd = ParsedDoc()
    assert pd.text == ""
    assert pd.pages == []
    assert pd.tables == []
    assert pd.metadata == {}
    assert pd.error is None
    assert pd.num_pages == 0
    assert pd.num_chars == 0


def test_parsed_doc_num_pages_falls_back_to_metadata():
    """If pages is empty but metadata.num_pages is set (PDF metadata), use that."""
    pd = ParsedDoc(metadata={"num_pages": 17})
    assert pd.num_pages == 17


def test_parse_document_never_raises_on_garbage(tmp_path):
    """Even on a corrupt file, parse_document returns an error rather
    than raising — research pipelines depend on this."""
    src = tmp_path / "garbage.pdf"
    src.write_bytes(b"this is not a real PDF \x00\x01\x02\xff\xfe")
    doc = parse_document(src)
    assert isinstance(doc, ParsedDoc)
    # error may or may not be set depending on the underlying lib's
    # tolerance — what matters is that we got a ParsedDoc back.


def test_parse_document_invalid_max_pages_uses_default(tmp_path):
    """max_pages=None / 0 / negative → fall through to DEFAULT_MAX_PAGES."""
    p = tmp_path / "x.txt"
    p.write_text("hi")
    for bad in (None, 0, -5):
        doc = parse_document(p, max_pages=bad)
        assert doc.error is None
